"""Extract text from user-uploaded files (Telegram documents)."""

from __future__ import annotations

import csv
import io
import zipfile
from dataclasses import dataclass
from xml.etree import ElementTree

_DEFAULT_MAX_BYTES = 15 * 1024 * 1024
_DEFAULT_MAX_CHARS = 48_000


def _limits() -> tuple[int, int]:
    try:
        from personal_ai_os.config import get_settings

        s = get_settings()
        return s.file_upload_max_bytes, s.file_extract_max_chars
    except Exception:
        return _DEFAULT_MAX_BYTES, _DEFAULT_MAX_CHARS


@dataclass
class IngestResult:
    filename: str
    kind: str
    text: str
    truncated: bool = False
    error: str | None = None


def _clip(text: str) -> tuple[str, bool]:
    _, limit = _limits()
    if len(text) <= limit:
        return text, False
    return text[:limit] + "\n\n…[текст обрезан по лимиту]", True


def _parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _parse_csv(data: bytes) -> str:
    text = _parse_txt(data)
    reader = csv.reader(io.StringIO(text))
    rows: list[str] = []
    for i, row in enumerate(reader):
        if i >= 500:
            rows.append("…[строки обрезаны]")
            break
        rows.append(" | ".join(row))
    return "\n".join(rows)


def _parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets[:10]:
        parts.append(f"### Лист: {sheet.title}")
        row_count = 0
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if not any(cells):
                continue
            parts.append("\t".join(cells))
            row_count += 1
            if row_count >= 500:
                parts.append("…[строки листа обрезаны]")
                break
    wb.close()
    return "\n".join(parts)


def _parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables[:20]:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for i, page in enumerate(reader.pages[:50]):
        t = page.extract_text() or ""
        if t.strip():
            parts.append(f"--- стр. {i + 1} ---\n{t}")
    return "\n".join(parts)


def _parse_pptx(data: bytes) -> str:
    """Minimal PPTX: slide text from XML inside the zip."""
    parts: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = sorted(n for n in zf.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
        for idx, name in enumerate(names[:40], start=1):
            root = ElementTree.fromstring(zf.read(name))
            texts: list[str] = []
            for el in root.iter():
                if el.tag.endswith("}t") and el.text and el.text.strip():
                    texts.append(el.text.strip())
            if texts:
                parts.append(f"--- слайд {idx} ---\n" + "\n".join(texts))
    return "\n".join(parts)


def ingest_file(data: bytes, filename: str) -> IngestResult:
    """Parse file bytes to plain text for LLM context."""
    max_bytes, _ = _limits()
    if len(data) > max_bytes:
        return IngestResult(
            filename=filename,
            kind="unknown",
            text="",
            error=f"Файл больше {max_bytes // (1024 * 1024)} МБ.",
        )

    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "").strip()
    try:
        if ext in ("txt", "md", "markdown", "log", "json", "yaml", "yml"):
            raw = _parse_txt(data)
            kind = "text"
        elif ext == "csv":
            raw = _parse_csv(data)
            kind = "csv"
        elif ext in ("xlsx", "xlsm"):
            raw = _parse_xlsx(data)
            kind = "excel"
        elif ext == "docx":
            raw = _parse_docx(data)
            kind = "word"
        elif ext == "pdf":
            raw = _parse_pdf(data)
            kind = "pdf"
        elif ext == "pptx":
            raw = _parse_pptx(data)
            kind = "pptx"
        elif ext == "doc":
            return IngestResult(
                filename=filename,
                kind="word",
                text="",
                error="Старый .doc не поддерживается — сохрани как .docx.",
            )
        elif ext == "xls":
            return IngestResult(
                filename=filename,
                kind="excel",
                text="",
                error="Старый .xls не поддерживается — сохрани как .xlsx.",
            )
        else:
            return IngestResult(
                filename=filename,
                kind="unknown",
                text="",
                error=f"Формат .{ext or '?'} пока не поддерживается.",
            )
    except Exception as exc:
        return IngestResult(
            filename=filename,
            kind=ext or "unknown",
            text="",
            error=f"Не удалось разобрать файл: {exc}",
        )

    text, truncated = _clip(raw.strip())
    return IngestResult(
        filename=filename,
        kind=kind,
        text=text,
        truncated=truncated,
    )


def format_ingest_for_prompt(result: IngestResult, *, user_caption: str = "") -> str:
    """Turn ingest result into a user-message appendix for agents."""
    lines = [f"[Вложение: {result.filename}, тип: {result.kind}]"]
    if result.error:
        lines.append(f"Ошибка разбора: {result.error}")
    elif result.text:
        lines.append(result.text)
    else:
        lines.append("(пустой или без извлекаемого текста)")
    if user_caption.strip():
        lines.append(f"\nПодпись пользователя: {user_caption.strip()}")
    return "\n".join(lines)
